# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Configurar margenes ──────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Estilo base ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helpers ──────────────────────────────────────────────────────────────────
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_h1(text):
    return doc.add_heading(text, level=1)

def add_h2(text):
    return doc.add_heading(text, level=2)

def add_h3(text):
    return doc.add_heading(text, level=3)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(bold_prefix, normal_text):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(normal_text)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # linea horizontal
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()  # espacio
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('INFORME TECNICO')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Estrategias de Pruebas, Pipeline CI/CD\ny Validacion del Software')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('E-Commerce "Lacteos Aqui Es Purita"')
r3.font.size = Pt(14)
r3.bold = True

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Asignatura: Desarrollo de Software III\n').font.size = Pt(11)
info.add_run('Universidad del Valle - 2025\n').font.size = Pt(11)

doc.add_paragraph()

equipo = doc.add_paragraph()
equipo.alignment = WD_ALIGN_PARAGRAPH.CENTER
equipo.add_run('Integrantes:\n').bold = True
equipo.add_run('Kevin Camilo Molina Salazar - Frontend\n')
equipo.add_run('Santiago Bedon Gomez - Backend\n')
equipo.add_run('Daniela Martinez Fontal - Admin BD y Documentacion\n')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECCION 1: PRUEBAS DE AISLAMIENTO
# ════════════════════════════════════════════════════════════════════════════
add_h1('1. Pruebas de Aislamiento, Arquitectura y Microservicios')

add_body(
    'Como parte del rigor tecnico aplicado al proyecto, se disenaron procesos especificos '
    'para evaluar el desacoplamiento real entre las capas del sistema y el comportamiento del '
    'microservicio de autenticacion de forma completamente independiente al framework Django. '
    'Estas pruebas garantizan que la Arquitectura Hexagonal implementada cumple su promesa '
    'de independencia entre capas.'
)

add_h2('1.1 Aislamiento Funcional (test_run.py)')

add_body(
    'Para validar que la Capa de Dominio no esta acoplada al framework web (Django) ni a un '
    'motor de base de datos especifico, se construyo un script de testeo nativo llamado '
    'test_run.py. Este script simula exactamente lo que haria un API Gateway real:'
)

add_bullet('', 'Construye un "evento" con un cuerpo JSON que contiene username y password.')
add_bullet('', 'Lo lanza directamente contra la funcion lambda_handler() del microservicio.')
add_bullet('', 'Ejecuta dos escenarios controlados:')

p_sub = doc.add_paragraph(style='List Bullet 2')
r_s = p_sub.add_run('Credenciales correctas ')
r_s2 = p_sub.add_run('(esperando HTTP 200 OK)')
r_s2.bold = True

p_sub2 = doc.add_paragraph(style='List Bullet 2')
r_s3 = p_sub2.add_run('Contrasena incorrecta ')
r_s4 = p_sub2.add_run('(esperando HTTP 401 Unauthorized)')
r_s4.bold = True

add_body(
    'Esto verifica que la logica de negocio responde correctamente sin necesidad de levantar '
    'el servidor Django ni conectar a ninguna base de datos real.'
)

add_h2('1.2 Repositorio en Memoria Volatil (in_memory_repository.py)')

add_body(
    'Durante las pruebas del microservicio, el sistema es alimentado por un '
    'InMemoryUserRepository, un adaptador de persistencia completamente simulado en '
    'memoria RAM. Este repositorio:'
)

add_bullet('', 'Carga datos de usuarios de prueba directamente en un diccionario Python.')
add_bullet('', 'Reemplaza perfectamente al adaptador real de base de datos.')
add_bullet('', 'No requiere cambiar una sola linea del nucleo de negocio.')

add_body(
    'Esto cumple la promesa fundamental de la Arquitectura Hexagonal: el dominio es '
    'completamente agnostico a la infraestructura de persistencia. Puede funcionar con '
    'SQLite, PostgreSQL o un simple diccionario en memoria, sin modificar el codigo del negocio.'
)

add_h2('1.3 Gestion Multi-Entorno Segura (os.environ)')

add_body(
    'Se diseno una estrategia de configuracion dinamica mediante variables de entorno '
    '(os.environ["AUTH_MICROSERVICE_URL"]) que le indica a la aplicacion si debe operar '
    'hacia direcciones locales (http://localhost:8001) durante el desarrollo, o hacia el '
    'proveedor PaaS configurado en la nube durante produccion. Esta distincion garantiza que '
    'ningun desarrollador del equipo pueda accidentalmente conectar codigo local a la base '
    'de datos de produccion.'
)

add_separator()

# ════════════════════════════════════════════════════════════════════════════
# SECCION 2: PIPELINE CI/CD
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1('2. Pipeline CI/CD Automatizado con GitHub Actions')

add_body(
    'Ningun cambio es promovido a la plataforma principal de manera manual ni arbitraria. '
    'El entorno automatizado, orquestado en el archivo .github/workflows/pipeline-ci-cd.yml '
    'bajo el nombre "Pipeline E-Commerce Purita", actua como un filtro estricto de calidad.'
)

add_h2('2.1 Plataforma utilizada: GitHub Actions')

add_body(
    'GitHub Actions es la plataforma de integracion y entrega continua integrada directamente '
    'en GitHub. Permite definir flujos de trabajo (workflows) en archivos YAML que se ejecutan '
    'automaticamente en respuesta a eventos del repositorio (como un push a la rama main). Fue '
    'elegida por su integracion nativa con el repositorio, runners gratuitos (ubuntu-latest) y '
    'su ecosistema de actions reutilizables.'
)

add_body(
    'El pipeline se compone de 5 fases tecnicas secuenciales que se disparan unicamente '
    'con cada evento push en la rama main:'
)

add_h2('2.2 Los 5 Pasos del Pipeline')

# Paso 1
add_h3('Paso 1 - Descargar y Aislar (Checkout)')
add_body(
    'Se instancia una maquina virtual inmaculada basada en ubuntu-latest dentro de la '
    'infraestructura de GitHub. A traves de la accion oficial actions/checkout@v3, se efectua '
    'una clonacion limpia y completa del repositorio, garantizando que el runner trabaje siempre '
    'con el codigo mas reciente sin archivos residuales de ejecuciones anteriores.'
)

# Paso 2
add_h3('Paso 2 - Compilar Entorno (Instalar Python y Dependencias)')
add_body(
    'Mediante actions/setup-python@v4 se instala y configura el runtime exacto de Python 3.11 '
    'en la maquina virtual. Inmediatamente despues, pip evalua e instala de forma determinista '
    'todas las librerias declaradas en requirements.txt:'
)
add_bullet('', 'Django 5.1.7 (framework principal)')
add_bullet('', 'Pillow (procesamiento de imagenes)')
add_bullet('', 'psycopg2-binary (conexion a PostgreSQL)')
add_bullet('', 'ReportLab (generacion de PDFs)')
add_bullet('', 'openpyxl (generacion de Excel)')
add_bullet('', 'cloudinary / django-cloudinary-storage (CDN de imagenes)')
add_bullet('', 'whitenoise (archivos estaticos)')
add_bullet('', 'qrcode (generacion de codigos QR)')
add_body(
    'Esto asegura que el entorno de pruebas sea identico al de produccion.'
)

# Paso 3
add_h3('Paso 3 - Pruebas e Integridad del Sistema (Django Checks y Migraciones)')
add_body(
    'Es el paso mas critico del pipeline y el corazon del control de calidad. GitHub Actions '
    'ejecuta dos comandos de validacion:'
)
add_bullet('python manage.py check: ', 
           'Analiza toda la configuracion de Django, modelos ORM, relaciones entre tablas y '
           'registros de las aplicaciones buscando inconsistencias o errores de configuracion.')
add_bullet('python manage.py makemigrations --dry-run: ', 
           'Simula la creacion de migraciones sin escribirlas, detectando si algun modelo fue '
           'modificado sin generar su respectiva migracion.')
add_body(
    'Si cualquiera de estos comandos falla, el pipeline se detiene completamente y el despliegue '
    'es abortado, protegiendo la base de datos PostgreSQL de produccion de estados corruptos.'
)

# Paso 4
add_h3('Paso 4 - Generar Version de Produccion (Collectstatic)')
add_body(
    'Se ejecuta "python manage.py collectstatic --noinput" para recolectar y centralizar todos '
    'los archivos estaticos del proyecto (CSS, JavaScript, imagenes del frontend) en el directorio '
    'staticfiles/. Este proceso prepara la version de produccion del frontend, permitiendo que '
    'WhiteNoise sirva estos recursos de forma ultrarapida y eficiente sin necesidad de un '
    'servidor web adicional como Nginx.'
)

# Paso 5
add_h3('Paso 5 - Despliegue Multi-Cloud Zero-Downtime (Publish)')
add_body(
    'Solo si los cuatro pasos anteriores culminan con exito al 100%, el codigo es promovido '
    'a la nube publica. El sistema cuenta con configuracion dual:'
)
add_bullet('render.yaml: ',
           'Define el servicio completo en Render.com: instalacion de dependencias, '
           'collectstatic, migrate y arranque con gunicorn login.wsgi:application.')
add_bullet('vercel.json: ',
           'Ofrece la alternativa serverless en Vercel, activable en el Paso 5 del pipeline '
           'mediante npx vercel --prod --token $VERCEL_TOKEN.')
add_body(
    'Esta arquitectura multi-cloud garantiza Zero-Downtime durante el despliegue: los usuarios '
    'que esten navegando el catalogo o finalizando una compra no experimentan ninguna interrupcion '
    'durante la publicacion de nuevas versiones.'
)

add_separator()

# ════════════════════════════════════════════════════════════════════════════
# SECCION 3: CONFIGURACION DE RENDER
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1('3. Configuracion de Render (render.yaml)')

add_body(
    'El archivo render.yaml define el servicio de hosting en Render.com con las siguientes '
    'instrucciones de build y arranque:'
)

add_h3('Build Command')
add_bullet('', 'pip install -r requirements.txt')
add_bullet('', 'python manage.py collectstatic --no-input')
add_bullet('', 'python manage.py migrate')

add_h3('Start Command')
add_bullet('', 'gunicorn login.wsgi:application')

add_h3('Variables de entorno gestionadas por Render')
add_bullet('SECRET_KEY: ', 'Generada automaticamente por Render.')
add_bullet('DATABASE_URL: ', 'Conexion a la base de datos PostgreSQL del mismo servicio.')
add_bullet('CLOUDINARY_CLOUD_NAME / API_KEY / API_SECRET: ', 'Credenciales para la CDN de imagenes.')
add_bullet('DEBUG=False: ', 'Modo produccion desactivando mensajes de depuracion.')
add_bullet('ALLOWED_HOSTS: ', '.onrender.com (dominio permitido).')

add_separator()

# ════════════════════════════════════════════════════════════════════════════
# SECCION 4: IMPORTANCIA
# ════════════════════════════════════════════════════════════════════════════
add_h1('4. Importancia del Pipeline en el Contexto Academico y Empresarial')

add_body(
    'La implementacion de este flujo de CI/CD no es unicamente un requisito tecnico; '
    'representa la adopcion de una cultura de ingenieria de software moderna y profesional. '
    'En un entorno empresarial real, cada minuto de caida o error en produccion representa '
    'perdidas economicas directas.'
)

add_body('El pipeline implementado garantiza cuatro propiedades fundamentales:')

add_bullet('Calidad garantizada: ',
           'Ningun codigo defectuoso llega jamas a los usuarios finales; el Paso 3 del pipeline '
           'actua como barrera impenetrable antes de cualquier despliegue.')
add_bullet('Velocidad de entrega: ',
           'El equipo puede publicar mejoras multiples veces al dia con total confianza, sin '
           'procedimientos manuales ni riesgos de errores humanos.')
add_bullet('Trazabilidad completa: ',
           'Cada ejecucion del pipeline queda registrada en GitHub con logs detallados, fechas '
           'y resultado de cada paso, permitiendo una auditoria completa del historial de despliegues.')
add_bullet('Automatizacion end-to-end: ',
           'GitHub Actions cierra el ciclo completo: desde el commit del desarrollador hasta '
           'el servidor en la nube de Render o Vercel, sin intervencion humana adicional.')

add_separator()

# ── PIE DE PAGINA ─────────────────────────────────────────────────────────────
doc.add_paragraph()
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_pie = pie.add_run('Universidad del Valle - Desarrollo de Software III - 2025')
r_pie.font.size = Pt(9)
r_pie.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── GUARDAR ───────────────────────────────────────────────────────────────────
out = 'EJEMPLO/Informe_Pipeline_y_Pruebas.docx'
doc.save(out)
print(f'Documento creado: {out}')
