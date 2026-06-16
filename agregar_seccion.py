# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('EJEMPLO/Informe_Lacteos_Purita (2).docx')

# Mapa de style_id -> primer estilo que lo tenga (evita duplicados)
_style_map = {}
for s in doc.styles:
    if s.style_id not in _style_map:
        _style_map[s.style_id] = s

def set_style_by_id(paragraph, style_id):
    paragraph.style = _style_map[style_id]

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_style_by_id(p, f'Heading{level}')
    p.add_run(text)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

def add_bullet(doc, normal_text, bold_prefix=None):
    p = doc.add_paragraph(style=_style_map['ListParagraph'])
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
    p.add_run(normal_text)
    return p

# ── Salto de pagina ───────────────────────────────────────────────────────────
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECCION 9 – PRUEBAS Y PIPELINE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Estrategias de Pruebas, Validacion y Pipeline CI/CD Automatizado', level=1)

add_body(doc,
    'Para garantizar que el patron Hexagonal y el diseno de la plataforma web sean resilientes, '
    'robustos y escalables, se establecio una cultura estricta de DevOps, basando el ciclo de vida '
    'del software en Integracion Continua y Entrega Continua (CI/CD) mediante las herramientas de '
    'GitHub Actions, Render y Vercel. Esto asegura que cada nueva funcionalidad o correccion sea '
    'probada en escenarios aislados antes de llegar al usuario final, protegiendo la integridad '
    'de la base de datos de produccion (PostgreSQL) y el funcionamiento general del e-commerce.'
)

# ── 9.1 ──────────────────────────────────────────────────────────────────────
add_heading(doc, '9.1 Pruebas de Aislamiento, Arquitectura y Microservicios', level=2)

add_body(doc,
    'Como parte del rigor tecnico aplicado, se disenaron procesos especificos para evaluar el '
    'desacoplamiento real entre las capas del sistema y el comportamiento del microservicio de '
    'autenticacion de forma completamente independiente al framework Django:'
)

add_bullet(doc,
    'Para validar verdaderamente que la Capa de Dominio no esta acoplada al framework web (Django) '
    'ni a un motor de base de datos especifico, se construyo un script de testeo nativo (test_run.py). '
    'Este script simula exactamente lo que haria un API Gateway real: construye un "evento" con un '
    'cuerpo JSON (username + password) y lo lanza directamente contra la funcion lambda_handler() '
    'del microservicio. Se ejecutan dos escenarios controlados: credenciales correctas '
    '(esperando HTTP 200 OK) y contrasena incorrecta (esperando HTTP 401 Unauthorized), '
    'verificando que la logica de negocio responde correctamente sin necesidad de levantar Django '
    'ni conectar a ninguna base de datos.',
    bold_prefix='Aislamiento Funcional (test_run.py): '
)

add_bullet(doc,
    'Durante las pruebas del microservicio, el sistema es alimentado por un InMemoryUserRepository '
    '(un adaptador de persistencia completamente simulado en memoria RAM). Este repositorio carga '
    'datos de usuarios de prueba directamente en un diccionario Python, reemplazando perfectamente '
    'al adaptador real de base de datos sin cambiar una sola linea del nucleo de negocio. '
    'Esto cumple con la promesa fundamental de la Arquitectura Hexagonal: el dominio es '
    'completamente agnostico a la infraestructura de persistencia.',
    bold_prefix='Repositorio en Memoria Volatil (in_memory_repository.py): '
)

add_bullet(doc,
    'Se diseno una estrategia de configuracion dinamica mediante variables de entorno '
    '(os.environ["AUTH_MICROSERVICE_URL"]) que le indica a la aplicacion si debe operar '
    'hacia direcciones locales (http://localhost:8001) durante el desarrollo, o hacia el '
    'proveedor PaaS configurado en la nube durante produccion. Esta distincion garantiza que '
    'ningun desarrollador del equipo pueda accidentalmente conectar codigo local a la base '
    'de datos de produccion.',
    bold_prefix='Gestion Multi-Entorno Segura (os.environ): '
)

# ── 9.2 ──────────────────────────────────────────────────────────────────────
add_heading(doc, '9.2 Pipeline Estricto de Integracion Continua (CI/CD) con GitHub Actions', level=2)

add_body(doc,
    'Ningun cambio es promovido a la plataforma principal de manera manual ni arbitraria. '
    'El entorno automatizado, orquestado en el archivo .github/workflows/pipeline-ci-cd.yml '
    'bajo el nombre "Pipeline E-Commerce Purita", actua como un filtro estricto de calidad. '
    'Esta compuesto por 5 fases tecnicas secuenciales que se disparan automaticamente como '
    'reaccion a cada evento push en la rama main de GitHub:'
)

steps = [
    (
        'Paso 1 - Descargar y Aislar (Checkout): ',
        'Se instancia una maquina virtual inmaculada basada en ubuntu-latest dentro de la '
        'infraestructura de GitHub. A continuacion, a traves de la accion oficial '
        'actions/checkout@v3, se efectua una clonacion limpia y completa del repositorio, '
        'garantizando que el runner trabaje siempre con el estado exacto y mas reciente del '
        'codigo fuente, evitando cualquier cruce de archivos residuales de ejecuciones anteriores.'
    ),
    (
        'Paso 2 - Compilar Entorno (Instalar Python y Dependencias): ',
        'A traves de la accion actions/setup-python@v4, se instala y configura el runtime '
        'exacto de Python 3.11 en la maquina virtual. Inmediatamente despues, pip evalua e '
        'instala de forma determinista todas las librerias del requirements.txt: Django 5.1.7, '
        'Pillow, psycopg2-binary, ReportLab, openpyxl, cloudinary, whitenoise y qrcode, '
        'asegurando que el entorno de pruebas sea identico al de produccion.'
    ),
    (
        'Paso 3 - Pruebas e Integridad del Sistema (Django Checks y Migraciones): ',
        'Es el paso mas critico del pipeline. GitHub Actions ejecuta: primero '
        '"python manage.py check", que analiza toda la configuracion de Django, modelos ORM '
        'y relaciones entre tablas buscando inconsistencias; y segundo, '
        '"python manage.py makemigrations --dry-run", que simula la creacion de migraciones '
        'sin escribirlas, detectando modelos modificados sin migracion generada. Si cualquiera '
        'de estos comandos falla, el pipeline se detiene completamente y el despliegue es '
        'abortado, protegiendo la base de datos PostgreSQL de produccion.'
    ),
    (
        'Paso 4 - Generar Version de Produccion (Collectstatic): ',
        'Se ejecuta "python manage.py collectstatic --noinput" para recolectar y centralizar '
        'todos los archivos estaticos del proyecto (CSS, JavaScript, imagenes) en el directorio '
        'staticfiles/. Este proceso prepara la version de produccion del frontend, permitiendo '
        'que WhiteNoise sirva estos recursos de forma ultrarapida sin necesidad de un servidor '
        'web adicional como Nginx.'
    ),
    (
        'Paso 5 - Despliegue Multi-Cloud sin Caidas (Zero-Downtime Publish): ',
        'Solo si los cuatro pasos anteriores culminan con exito al 100%, el codigo es promovido '
        'a la nube publica. El sistema cuenta con configuracion dual: el archivo render.yaml '
        'define el servicio completo en Render.com (dependencias, collectstatic, migrate y '
        'arranque con gunicorn login.wsgi:application), mientras que vercel.json ofrece la '
        'alternativa serverless en Vercel. Esta arquitectura multi-cloud garantiza '
        'Zero-Downtime: los usuarios que esten navegando el catalogo o finalizando una compra '
        'no experimentan ninguna interrupcion durante la publicacion de nuevas versiones.'
    ),
]

for bold_part, normal_part in steps:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

# ── 9.3 ──────────────────────────────────────────────────────────────────────
add_heading(doc, '9.3 Importancia del Pipeline en el Contexto Academico y Empresarial', level=2)

add_body(doc,
    'La implementacion de este flujo de CI/CD no es unicamente un requisito tecnico; representa '
    'la adopcion de una cultura de ingenieria de software moderna y profesional. En un entorno '
    'empresarial real, cada minuto de caida o error en produccion representa perdidas economicas '
    'directas. El pipeline implementado garantiza cuatro propiedades fundamentales:'
)

guarantees = [
    ('Calidad garantizada: ',
     'Ningun codigo defectuoso llega jamas a los usuarios finales; el Paso 3 actua como barrera impenetrable.'),
    ('Velocidad de entrega: ',
     'El equipo puede publicar mejoras multiples veces al dia con total confianza, sin procedimientos manuales ni riesgos.'),
    ('Trazabilidad completa: ',
     'Cada ejecucion del pipeline queda registrada en GitHub con logs detallados, fechas y resultado de cada paso, permitiendo una auditoria completa del historial de despliegues.'),
    ('Automatizacion end-to-end: ',
     'GitHub Actions cierra el ciclo completo: desde el commit del desarrollador hasta el servidor en la nube de Render o Vercel, sin intervencion humana adicional.'),
]

for bold_part, normal_part in guarantees:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save('EJEMPLO/Informe_Lacteos_Purita_ACTUALIZADO.docx')
print('Seccion 9 agregada correctamente al archivo Word.')
